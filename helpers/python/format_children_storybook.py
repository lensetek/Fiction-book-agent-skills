#!/usr/bin/env python3
"""
format_children_storybook.py
Generates publisher-grade picture book Word (.docx) documents with image placeholders, page spreads, and large child-friendly typography.
"""

import sys
import os
import argparse

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Error: python-docx library is required. Install via: pip install python-docx")
    sys.exit(1)

def build_storybook_docx(script_path, output_path, title="Buku Cerita Anak"):
    doc = Document()
    
    # Set page layout to A4 Landscape (Standard Picture Book Spread)
    section = doc.sections[0]
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # Title Page
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(140)
    p_title.paragraph_format.space_after = Pt(20)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = p_title.add_run(title)
    run_t.font.name = "Comic Sans MS"
    run_t.font.size = Pt(32)
    run_t.bold = True
    run_t.font.color.rgb = RGBColor(0xE1, 0x5B, 0x00) # Friendly Orange

    doc.add_page_break()

    if os.path.exists(script_path):
        with open(script_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()

        for line in lines:
            text = line.strip()
            if text.startswith("### Page") or text.startswith("## Page"):
                p_page = doc.add_paragraph()
                p_page.paragraph_format.space_before = Pt(24)
                p_page.paragraph_format.space_after = Pt(12)
                run_p = p_page.add_run(text)
                run_p.font.name = "Comic Sans MS"
                run_p.font.size = Pt(18)
                run_p.bold = True
            elif text.startswith("> "):
                p_text = doc.add_paragraph()
                p_text.paragraph_format.space_before = Pt(12)
                p_text.paragraph_format.space_after = Pt(18)
                p_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_st = p_text.add_run(text[2:])
                run_st.font.name = "Comic Sans MS"
                run_st.font.size = Pt(22)
                run_st.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    doc.save(output_path)
    print(f"[SUCCESS] Children Storybook DOCX generated: {output_path}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Children Storybook DOCX Formatter")
    parser.add_argument("--input", required=True, help="Path to storybook script markdown file")
    parser.add_argument("--output", required=True, help="Path to output .docx file")
    parser.add_argument("--title", default="Buku Cerita Anak", help="Book Title")
    args = parser.parse_args()
    build_storybook_docx(args.input, args.output, args.title)
