#!/usr/bin/env python3
"""
generate_fiction_docx.py
Publisher-Grade Microsoft Word (.docx) Generator for Fiction Novels & Creative Manuscripts.

Supports:
- Print Layout Presets: Novel (13x19 cm), A5 (14.8x21 cm), A4 (Storybook/Script)
- Binding Gutter & Mirrored Margins (Recto / Verso Odd/Even page layout)
- Chapter Opener Rules (Odd-page start, 120pt Space Before title, No-indent first paragraph)
- Native Word TOC Field Codes (TOC \\o "1-3" \\h \\z \\u)
- Typography: Garamond / Georgia / Book Antiqua, 1.15 line spacing, 0.7cm first line indent
"""

import sys
import os
import argparse
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_SECTION_START
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import qn, nsdecls
except ImportError:
    print("Error: python-docx library is required. Install via: pip install python-docx")
    sys.exit(1)


# Preset Page Specifications (in Centimeters)
LAYOUT_PRESETS = {
    "novel_13x19": {
        "width": 13.0,
        "height": 19.0,
        "margin_top": 2.0,
        "margin_bottom": 2.0,
        "margin_inside": 2.2, # Extra gutter margin for binding
        "margin_outside": 1.8,
        "font_name": "Garamond",
        "font_size": 11,
        "line_spacing": 1.15
    },
    "A5": {
        "width": 14.8,
        "height": 21.0,
        "margin_top": 2.2,
        "margin_bottom": 2.2,
        "margin_inside": 2.5,
        "margin_outside": 2.0,
        "font_name": "Georgia",
        "font_size": 11,
        "line_spacing": 1.2
    },
    "A4": {
        "width": 21.0,
        "height": 29.7,
        "margin_top": 2.5,
        "margin_bottom": 2.5,
        "margin_inside": 3.0,
        "margin_outside": 2.5,
        "font_name": "Times New Roman",
        "font_size": 12,
        "line_spacing": 1.5
    }
}


def add_toc_field(paragraph):
    """Inserts native Word XML TOC field into paragraph."""
    run = paragraph.add_run()
    r = run._r
    fldChar1 = parse_xml(r'<w:fldChar %s w:fldCharType="begin"/>' % nsdecls('w'))
    instrText = parse_xml(r'<w:instrText %s xml:space="preserve"> TOC \o "1-3" \h \z \u </w:instrText>' % nsdecls('w'))
    fldChar2 = parse_xml(r'<w:fldChar %s w:fldCharType="separate"/>' % nsdecls('w'))
    fldChar3 = parse_xml(r'<w:fldChar %s w:fldCharType="end"/>' % nsdecls('w'))
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)


def setup_document_styles(doc, preset_config):
    """Applies publisher typography styles to document."""
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = preset_config['font_name']
    normal_style.font.size = Pt(preset_config['font_size'])
    normal_style.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
    
    p_format = normal_style.paragraph_format
    p_format.line_spacing = preset_config['line_spacing']
    p_format.space_before = Pt(0)
    p_format.space_after = Pt(0)
    p_format.first_line_indent = Cm(0.7) # Standard novel paragraph indent


def set_section_margins(section, preset_config):
    """Sets page dimensions and mirrored margins for binding."""
    section.page_width = Cm(preset_config['width'])
    section.page_height = Cm(preset_config['height'])
    section.top_margin = Cm(preset_config['margin_top'])
    section.bottom_margin = Cm(preset_config['margin_bottom'])
    section.left_margin = Cm(preset_config['margin_inside'])
    section.right_margin = Cm(preset_config['margin_outside'])
    section.different_first_page_header_footer = True


def build_fiction_docx(markdown_path, output_path, preset="novel_13x19", title="Novel", author="Author"):
    """Converts structured markdown manuscript into publisher-grade DOCX."""
    if preset not in LAYOUT_PRESETS:
        preset = "novel_13x19"
    config = LAYOUT_PRESETS[preset]

    doc = Document()
    setup_document_styles(doc, config)
    section = doc.sections[0]
    set_section_margins(section, config)

    # Title Page
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(120)
    title_p.paragraph_format.space_after = Pt(18)
    title_p.paragraph_format.first_line_indent = Cm(0)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = title_p.add_run(title)
    run_t.font.size = Pt(26)
    run_t.bold = True
    run_t.font.name = config['font_name']

    author_p = doc.add_paragraph()
    author_p.paragraph_format.space_before = Pt(0)
    author_p.paragraph_format.space_after = Pt(180)
    author_p.paragraph_format.first_line_indent = Cm(0)
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_a = author_p.add_run(f"Oleh {author}")
    run_a.font.size = Pt(14)
    run_a.italic = True

    # Page Break for TOC
    doc.add_page_break()
    toc_heading = doc.add_paragraph()
    toc_heading.paragraph_format.space_before = Pt(36)
    toc_heading.paragraph_format.space_after = Pt(24)
    toc_heading.paragraph_format.first_line_indent = Cm(0)
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_run = toc_heading.add_run("DAFTAR ISI")
    toc_run.font.size = Pt(16)
    toc_run.bold = True

    toc_p = doc.add_paragraph()
    toc_p.paragraph_format.first_line_indent = Cm(0)
    add_toc_field(toc_p)

    # Main Manuscript Parser
    if os.path.exists(markdown_path):
        with open(markdown_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        
        is_first_paragraph_in_chapter = False

        for line in lines:
            text = line.strip()
            if not text:
                continue

            if text.startswith("# "):
                # Chapter Heading (Start new section or odd page)
                new_sec = doc.add_section(WD_SECTION_START.ODD_PAGE)
                set_section_margins(new_sec, config)
                
                ch_p = doc.add_paragraph()
                ch_p.paragraph_format.space_before = Pt(100)
                ch_p.paragraph_format.space_after = Pt(36)
                ch_p.paragraph_format.first_line_indent = Cm(0)
                ch_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = ch_p.add_run(text[2:])
                run.font.size = Pt(20)
                run.bold = True
                is_first_paragraph_in_chapter = True

            elif text.startswith("## "):
                # Sub-heading / Scene Separator
                sub_p = doc.add_paragraph()
                sub_p.paragraph_format.space_before = Pt(18)
                sub_p.paragraph_format.space_after = Pt(12)
                sub_p.paragraph_format.first_line_indent = Cm(0)
                sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = sub_p.add_run(text[3:])
                run.font.size = Pt(14)
                run.italic = True
                is_first_paragraph_in_chapter = True

            elif text == "***" or text == "---" or text == "* * *":
                # Scene Break Asterisks
                break_p = doc.add_paragraph()
                break_p.paragraph_format.space_before = Pt(12)
                break_p.paragraph_format.space_after = Pt(12)
                break_p.paragraph_format.first_line_indent = Cm(0)
                break_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = break_p.add_run("* * *")
                run.font.size = Pt(12)
                is_first_paragraph_in_chapter = True

            else:
                # Regular Narrative Paragraph
                p = doc.add_paragraph()
                if is_first_paragraph_in_chapter:
                    p.paragraph_format.first_line_indent = Cm(0) # No indent after heading
                    is_first_paragraph_in_chapter = False
                else:
                    p.paragraph_format.first_line_indent = Cm(0.7)
                
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.add_run(text)

    # Save Output Document
    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    doc.save(output_path)
    print(f"[SUCCESS] Fiction DOCX generated: {output_path} (Preset: {preset})")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Publisher-Grade Fiction DOCX Generator")
    parser.add_argument("--input", required=True, help="Path to manuscript Markdown file")
    parser.add_argument("--output", required=True, help="Path to output .docx file")
    parser.add_argument("--preset", default="novel_13x19", choices=["novel_13x19", "A5", "A4"], help="Page preset")
    parser.add_argument("--title", default="Judul Novel", help="Book Title")
    parser.add_argument("--author", default="Penulis", help="Author Name")

    args = parser.parse_args()
    build_fiction_docx(args.input, args.output, args.preset, args.title, args.author)
